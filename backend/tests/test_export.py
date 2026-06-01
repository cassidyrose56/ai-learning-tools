import io
import re
import zipfile

import pytest
from docx import Document

from app.export import (
    render_docx,
    safe_filename,
    StoryInput,
)


def make_story(**over):
    base = dict(
        child_name="Maya",
        topic="Soccer",
        genre="fiction",
        text="Once upon a time. The end.",
        reading_level="3",
        pages=2,
        include_drawing_box=False,
    )
    base.update(over)
    return StoryInput(**base)


def test_docx_contains_title_and_body():
    blob = render_docx(make_story())
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    assert any("For Maya" in p and "Soccer" in p for p in paragraphs)
    assert any("Once upon a time" in p for p in paragraphs)


def test_docx_ignores_layout_fields_without_error():
    # include_drawing_box=True must not affect docx output
    blob = render_docx(make_story(include_drawing_box=True))
    Document(io.BytesIO(blob))  # parses without error


def test_safe_filename_replaces_unsafe_chars():
    assert safe_filename("Maya", "The Moon") == "Maya_The_Moon"
    assert safe_filename("Ma ya", "Soccer/Football") == "Ma_ya_Soccer_Football"
    assert safe_filename("Maya", "..hi..") == "Maya_hi"


def test_safe_filename_never_empty():
    assert safe_filename("", "") == "story_story"


from app.export import split_into_pages


def test_split_into_pages_returns_exact_chunk_count():
    text = "One two three. Four five. Six seven eight. Nine ten. Eleven twelve."
    chunks = split_into_pages(text, 3)
    assert len(chunks) == 3


def test_split_into_pages_never_breaks_mid_sentence():
    text = "Sentence one. Sentence two? Sentence three! Sentence four."
    chunks = split_into_pages(text, 2)
    for chunk in chunks:
        # each chunk must end with terminal punctuation
        assert chunk.rstrip()[-1] in ".!?"


def test_split_into_pages_groups_consecutive_sentences():
    text = " ".join(f"Sentence {i}." for i in range(1, 11))
    chunks = split_into_pages(text, 3)
    # concatenating chunks reconstructs the sentence sequence
    rejoined = " ".join(c.strip() for c in chunks)
    assert "Sentence 1." in rejoined
    assert "Sentence 10." in rejoined
    # chunks ordered
    indices = [int(re.search(r"Sentence (\d+)", c).group(1)) for c in chunks]
    assert indices == sorted(indices)


def test_split_into_pages_balanced_lengths():
    text = " ".join(f"Word{i} word{i}." for i in range(50))
    chunks = split_into_pages(text, 4)
    lens = [len(c.split()) for c in chunks]
    # roughly balanced - no chunk more than 2x the smallest
    assert max(lens) <= 2 * max(min(lens), 1)


def test_split_into_pages_handles_fewer_sentences_than_pages():
    text = "Only one sentence."
    chunks = split_into_pages(text, 3)
    assert len(chunks) == 3
    # first chunk has the content; rest may be empty
    assert chunks[0].strip().endswith(".")


from pypdf import PdfReader

from app.export import render_pdf
from app.pedagogy import FONT_SIZES, LINE_SPACING


def test_pdf_has_correct_page_count():
    story = make_story(pages=3)
    blob = render_pdf(story)
    reader = PdfReader(io.BytesIO(blob))
    assert len(reader.pages) == 3


def test_pdf_begins_with_pdf_signature():
    blob = render_pdf(make_story())
    assert blob[:4] == b"%PDF"
    assert len(blob) > 500


def test_pdf_draws_box_when_enabled(monkeypatch):
    # spy on canvas.rect calls
    rect_calls: list[tuple] = []
    from reportlab.pdfgen import canvas as canvas_mod

    real_rect = canvas_mod.Canvas.rect

    def spy_rect(self, *args, **kwargs):
        rect_calls.append((args, kwargs))
        return real_rect(self, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "rect", spy_rect)

    render_pdf(make_story(pages=2, include_drawing_box=True))
    box_with_box = len(rect_calls)
    rect_calls.clear()

    render_pdf(make_story(pages=2, include_drawing_box=False))
    box_without_box = len(rect_calls)

    # Box renders on page 1 only, regardless of how many pages the story has.
    assert box_with_box == 1
    assert box_without_box == 0


def test_pdf_font_size_per_reading_level():
    assert FONT_SIZES == {"K": 24, "1": 20, "2": 18, "3": 16, "4": 14, "5": 14}


def test_pdf_line_spacing_per_reading_level():
    assert LINE_SPACING == {
        "K": 1.8, "1": 1.7, "2": 1.6, "3": 1.55, "4": 1.5, "5": 1.5,
    }


def test_docx_title_uses_colon_not_em_dash():
    blob = render_docx(make_story(child_name="Maya", topic="Soccer"))
    doc = Document(io.BytesIO(blob))
    titles = [p.text for p in doc.paragraphs if p.text]
    assert any('For Maya: "Soccer"' in t for t in titles)
    assert not any("—" in t for t in titles)


def test_pdf_preserves_paragraph_breaks(monkeypatch):
    # Render a 1-page story whose text has three paragraphs. The drawn
    # y-coordinates must drop monotonically across body lines, and the
    # gap at each paragraph boundary must be strictly larger than the
    # gap inside a paragraph.
    from reportlab.pdfgen import canvas as canvas_mod

    calls: list[tuple[float, str]] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        calls.append((y, text))
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)

    para_a = "Alpha one. Alpha two. Alpha three."
    para_b = "Bravo one. Bravo two. Bravo three."
    para_c = "Charlie one. Charlie two. Charlie three."
    text = f"{para_a}\n\n{para_b}\n\n{para_c}"

    render_pdf(make_story(reading_level="3", pages=1, text=text))

    body = calls
    ys = [y for (y, _) in body]
    # monotonically decreasing
    assert all(ys[i] > ys[i + 1] for i in range(len(ys) - 1)), ys

    # Locate the first drawString that contains text from each paragraph.
    def first_idx(token: str) -> int:
        return next(i for i, (_, t) in enumerate(body) if token in t)

    a_first = first_idx("Alpha")
    b_first = first_idx("Bravo")
    c_first = first_idx("Charlie")

    # intra-paragraph leading: gap inside Alpha (if Alpha spans multiple lines
    # at this font size it will; if not, fall back to comparing the gap right
    # before each paragraph break instead).
    leading = ys[a_first] - ys[a_first + 1] if (b_first - a_first) > 1 else None

    gap_a_to_b = ys[b_first - 1] - ys[b_first]
    gap_b_to_c = ys[c_first - 1] - ys[c_first]

    if leading is not None:
        assert gap_a_to_b > leading, (gap_a_to_b, leading)
        assert gap_b_to_c > leading, (gap_b_to_c, leading)
    else:
        # Single-line paragraphs: the paragraph gaps must at least exceed the
        # body leading we know we are using.
        from app.pedagogy import FONT_SIZES, LINE_SPACING
        body_leading = FONT_SIZES["3"] * LINE_SPACING["3"]
        assert gap_a_to_b > body_leading
        assert gap_b_to_c > body_leading


def test_pdf_drawing_box_leaves_leading_times_1_5_gap(monkeypatch):
    # When include_drawing_box=True, the first body drawString's y should
    # equal (box_bottom - leading * 1.5). We compute box_bottom from the
    # module constants and the page size, then compare.
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas as canvas_mod
    from app.export import _MARGIN, _BOX_FRACTION
    from app.pedagogy import FONT_SIZES, LINE_SPACING

    width, height = LETTER
    font_size = FONT_SIZES["3"]
    leading = font_size * LINE_SPACING["3"]
    text_top_before_box = height - _MARGIN
    box_bottom = text_top_before_box - (height - 2 * _MARGIN) * _BOX_FRACTION
    expected_first_y = box_bottom - leading * 1.5

    drawn: list[tuple[float, str]] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        drawn.append((y, text))
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)
    render_pdf(make_story(reading_level="3", pages=1, include_drawing_box=True, text="Body."))

    body = drawn
    assert body, "expected at least one body line"
    first_y = body[0][0]
    # The first body line sits at or below (box_bottom - leading * 1.5).
    # With fill-scaling it may sit lower; never above.
    assert first_y <= expected_first_y + 0.01, (first_y, expected_first_y)


def test_pdf_uses_per_grade_leading(monkeypatch):
    # render the same long text at K vs grade 5; K's larger leading should
    # fit fewer lines per page, producing a smaller cursor_y per drawString.
    from reportlab.pdfgen import canvas as canvas_mod

    y_positions: list[float] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        y_positions.append(y)
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)

    long_text = " ".join(f"word{i}." for i in range(200))

    render_pdf(make_story(reading_level="K", pages=1, text=long_text))
    k_gaps = [y_positions[i] - y_positions[i + 1] for i in range(len(y_positions) - 1)]
    k_leading = max(k_gaps)  # title-to-body jump is larger; ignore via max-of-body
    y_positions.clear()

    render_pdf(make_story(reading_level="5", pages=1, text=long_text))
    g5_gaps = [y_positions[i] - y_positions[i + 1] for i in range(len(y_positions) - 1)]
    g5_leading = max(g5_gaps)

    # K leading = 24 * 1.6 = 38.4; grade-5 leading = 14 * 1.3 = 18.2.
    # Even after picking the max-gap (which includes the title->body jump
    # in both cases), K's body leading must dominate grade 5's.
    assert k_leading > g5_leading


from app.export import render_bundle


def test_bundle_zip_contains_one_file_per_story():
    stories = [make_story(topic="Soccer"), make_story(topic="The Moon")]
    blob = render_bundle(stories, fmt="docx")
    with zipfile.ZipFile(io.BytesIO(blob)) as zf:
        names = sorted(zf.namelist())
    assert names == ["Maya_Soccer.docx", "Maya_The_Moon.docx"]


def test_bundle_pdf_zip_files_are_pdfs():
    stories = [make_story(topic="Soccer", text="One. Two.")]
    blob = render_bundle(stories, fmt="pdf")
    with zipfile.ZipFile(io.BytesIO(blob)) as zf:
        with zf.open("Maya_Soccer.pdf") as f:
            assert f.read(4) == b"%PDF"


def test_split_into_pages_preserves_closing_quotes():
    # All three sentences (including the one that ends with `."`) must
    # survive sentence tokenization and end up in the page output.
    text = (
        '"Hi," she said. "I am Maya. We just moved here." '
        '"Did you ride all the way in that truck?" Avi asked.'
    )
    chunks = split_into_pages(text, 1)
    page = chunks[0]
    assert '"Hi," she said.' in page
    assert '"I am Maya.' in page
    assert 'We just moved here."' in page
    assert '"Did you ride all the way in that truck?" Avi asked.' in page or (
        '"Did you ride all the way in that truck?"' in page
        and "Avi asked." in page
    )


def test_pdf_has_no_title(monkeypatch):
    from reportlab.pdfgen import canvas as canvas_mod

    drawn: list[str] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        drawn.append(text)
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)
    render_pdf(make_story(child_name="Maya", topic="Soccer", pages=1))
    assert not any(s.startswith("For ") for s in drawn)
    assert not any("Maya" in s for s in drawn)


def test_pdf_paragraphs_are_not_indented(monkeypatch):
    from reportlab.pdfgen import canvas as canvas_mod
    from app.export import _MARGIN

    calls: list[tuple[float, float, str]] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        calls.append((x, y, text))
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)

    para_a = "Alpha one. Alpha two. Alpha three. Alpha four."
    para_b = "Bravo one. Bravo two. Bravo three. Bravo four."
    text = f"{para_a}\n\n{para_b}"
    render_pdf(make_story(reading_level="3", pages=1, text=text))

    assert calls, "expected drawString calls"
    for x, _y, _t in calls:
        assert abs(x - _MARGIN) < 0.5, (x, _MARGIN)


def test_pdf_uses_natural_leading_for_sparse_content(monkeypatch):
    from reportlab.pdfgen import canvas as canvas_mod
    from reportlab.lib.pagesizes import LETTER
    from app.export import _MARGIN, _PARA_GAP_FRAC
    from app.pedagogy import FONT_SIZES, LINE_SPACING

    calls: list[tuple[float, str]] = []
    real_draw = canvas_mod.Canvas.drawString

    def spy_draw(self, x, y, text, *args, **kwargs):
        calls.append((y, text))
        return real_draw(self, x, y, text, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "drawString", spy_draw)

    # Three short sentences across two paragraphs; should NOT stretch.
    text = "Alpha one. Alpha two.\n\nBravo one."
    render_pdf(make_story(reading_level="3", pages=1, text=text))

    ys = [y for (y, _t) in calls]
    assert len(ys) >= 2, ys

    base_leading = FONT_SIZES["3"] * LINE_SPACING["3"]
    para_gap = base_leading * _PARA_GAP_FRAC
    deltas = [a - b for a, b in zip(ys, ys[1:])]

    # Every line-to-line drop is either exactly base_leading (intra-paragraph)
    # or base_leading + para_gap (across a paragraph boundary). The fill pass
    # is gone, so deltas never exceed those values.
    for d in deltas:
        assert (
            abs(d - base_leading) < 0.01
            or abs(d - (base_leading + para_gap)) < 0.01
        ), (d, base_leading, para_gap)

    # First body line sits near the top margin (no title above it).
    width, height = LETTER
    assert ys[0] > height - _MARGIN - base_leading * 3


def test_split_into_pages_balances_by_sentence_count():
    text = "One. Two. Three. Four. Five. Six. Seven. Eight."
    chunks = split_into_pages(text, 2)
    counts = [len(re.findall(r"[.!?]", c)) for c in chunks]
    assert counts == [4, 4], counts


def test_pdf_drawing_box_renders_only_on_first_page(monkeypatch):
    from reportlab.pdfgen import canvas as canvas_mod

    rect_calls: list[tuple] = []
    real_rect = canvas_mod.Canvas.rect

    def spy_rect(self, *args, **kwargs):
        rect_calls.append((args, kwargs))
        return real_rect(self, *args, **kwargs)

    monkeypatch.setattr(canvas_mod.Canvas, "rect", spy_rect)
    render_pdf(make_story(pages=3, include_drawing_box=True))
    assert len(rect_calls) == 1
