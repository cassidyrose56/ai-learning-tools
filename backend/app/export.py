import io
import re
import zipfile
from dataclasses import dataclass

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas as _canvas

from app.pedagogy import FONT_SIZES, LINE_SPACING


@dataclass
class StoryInput:
    child_name: str
    topic: str
    genre: str
    text: str
    reading_level: str
    pages: int
    include_drawing_box: bool


_UNSAFE = re.compile(r"[^A-Za-z0-9._-]+")


def safe_filename(child_name: str, topic: str) -> str:
    def clean(s: str) -> str:
        s = _UNSAFE.sub("_", s).strip("._-")
        return s or "story"
    return f"{clean(child_name)}_{clean(topic)}"


def render_docx(story: StoryInput) -> bytes:
    doc = Document()
    doc.add_heading(f'For {story.child_name}: "{story.topic}"', level=1)
    for para in [p.strip() for p in story.text.split("\n\n") if p.strip()]:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_SENTENCE_RE = re.compile(r"[^.!?]+[.!?]+[\"”’')\]]*(?:\s|$)")


def _tokenize_sentences(text: str) -> list[tuple[str, int]]:
    """Return [(sentence, paragraph_index), ...] preserving paragraph
    boundaries from the input's `\n\n` separators.
    """
    out: list[tuple[str, int]] = []
    for p_idx, para in enumerate(text.split("\n\n")):
        for s in _SENTENCE_RE.findall(para):
            s = s.strip()
            if s:
                out.append((s, p_idx))
    return out


def _join_sentences(items: list[tuple[str, int]]) -> str:
    """Join sentences with " " inside a paragraph and "\n\n" across
    paragraph boundaries, preserving the input paragraph structure.
    """
    if not items:
        return ""
    parts: list[str] = [items[0][0]]
    for prev, cur in zip(items, items[1:]):
        sep = " " if cur[1] == prev[1] else "\n\n"
        parts.append(sep + cur[0])
    return "".join(parts)


def split_into_pages(text: str, n: int) -> list[str]:
    if n <= 0:
        raise ValueError("n must be >= 1")
    sentences = _tokenize_sentences(text)
    if not sentences:
        return [""] * n
    if len(sentences) <= n:
        chunks = [_join_sentences([s]) for s in sentences]
        while len(chunks) < n:
            chunks.append("")
        return chunks

    total_words = sum(len(s[0].split()) for s in sentences)
    target = total_words / n
    chunks: list[str] = []
    current: list[tuple[str, int]] = []
    current_words = 0
    remaining_pages = n

    for i, item in enumerate(sentences):
        sent, _ = item
        words = len(sent.split())
        sentences_left = len(sentences) - i
        # ensure each remaining page gets at least one sentence
        if sentences_left <= remaining_pages - len(chunks) - 1 and current:
            chunks.append(_join_sentences(current))
            current = [item]
            current_words = words
            continue
        if current and current_words + words > target and len(chunks) < n - 1:
            chunks.append(_join_sentences(current))
            current = [item]
            current_words = words
        else:
            current.append(item)
            current_words += words

    if current:
        chunks.append(_join_sentences(current))
    while len(chunks) < n:
        chunks.append("")
    return chunks[:n]


_MARGIN = 0.75 * inch
_BOX_FRACTION = 0.45  # drawing box occupies top 45% of page
_BOX_GAP_LEADING = 1.5
_PARA_GAP_FRAC = 0.7  # paragraph gap as a fraction of leading
_FILL_MAX_SCALE = 2.0  # cap leading stretch to 2x baseline


def render_pdf(story: StoryInput) -> bytes:
    buf = io.BytesIO()
    width, height = LETTER
    c = _canvas.Canvas(buf, pagesize=LETTER)
    font_size = FONT_SIZES[story.reading_level]
    base_leading = font_size * LINE_SPACING[story.reading_level]

    chunks = split_into_pages(story.text, story.pages)
    text_x = _MARGIN
    text_width = width - 2 * _MARGIN

    for idx, chunk in enumerate(chunks):
        text_top = height - _MARGIN

        if story.include_drawing_box:
            box_top = text_top
            box_bottom = box_top - (height - 2 * _MARGIN) * _BOX_FRACTION
            box_height = box_top - box_bottom
            c.rect(_MARGIN, box_bottom, width - 2 * _MARGIN, box_height, stroke=1, fill=0)
            text_top = box_bottom - base_leading * _BOX_GAP_LEADING

        c.setFont("Helvetica", font_size)
        lines_per_para = _layout_chunk(chunk, text_width, font_size)

        available_height = text_top - _MARGIN
        scale = _compute_fill_scale(lines_per_para, available_height, base_leading)
        leading = base_leading * scale
        para_gap = base_leading * _PARA_GAP_FRAC * scale

        _draw_chunk(c, lines_per_para, text_x, text_top, leading, para_gap)

        if idx < len(chunks) - 1:
            c.showPage()

    c.showPage()  # close final page
    c.save()
    return buf.getvalue()


def _layout_paragraph(
    text: str, max_width: float, font_size: float
) -> list[str]:
    from reportlab.pdfbase.pdfmetrics import stringWidth

    font_name = "Helvetica"
    words = text.split()
    lines: list[str] = []
    current: list[str] = []

    def width_of(buf: list[str]) -> float:
        return stringWidth(" ".join(buf), font_name, font_size)

    for w in words:
        current.append(w)
        if width_of(current) > max_width:
            current.pop()
            if current:
                lines.append(" ".join(current))
            current = [w]
    if current:
        lines.append(" ".join(current))
    return lines


def _layout_chunk(
    text: str, max_width: float, font_size: float
) -> list[list[str]]:
    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    return [_layout_paragraph(p, max_width, font_size) for p in paragraphs]


def _compute_fill_scale(
    lines_per_para: list[list[str]],
    available_height: float,
    base_leading: float,
) -> float:
    # Last-baseline lands at: y_top - (N-1) * leading - para_breaks * gap.
    # We pick leading so that delta matches available_height.
    total_lines = sum(len(lp) for lp in lines_per_para)
    if total_lines <= 1:
        return 1.0
    para_breaks = max(0, len(lines_per_para) - 1)
    denom = (total_lines - 1) + para_breaks * _PARA_GAP_FRAC
    if denom <= 0:
        return 1.0
    raw_scale = available_height / (base_leading * denom)
    # Never shrink below baseline; never stretch beyond cap.
    return max(1.0, min(_FILL_MAX_SCALE, raw_scale))


def _draw_chunk(
    c,
    lines_per_para: list[list[str]],
    x: float,
    y_top: float,
    leading: float,
    para_gap: float,
) -> None:
    cursor_y = y_top
    for p_idx, lines in enumerate(lines_per_para):
        for line in lines:
            c.drawString(x, cursor_y, line)
            cursor_y -= leading
        if p_idx < len(lines_per_para) - 1:
            cursor_y -= para_gap


def render_bundle(stories: list[StoryInput], *, fmt: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for s in stories:
            base = safe_filename(s.child_name, s.topic)
            if fmt == "docx":
                zf.writestr(f"{base}.docx", render_docx(s))
            elif fmt == "pdf":
                zf.writestr(f"{base}.pdf", render_pdf(s))
            else:
                raise ValueError(f"unknown format {fmt!r}")
    return buf.getvalue()
